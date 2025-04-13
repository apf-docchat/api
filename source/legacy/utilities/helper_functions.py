import logging
import os, json, shutil, re
import pathlib
import uuid
from flask import request
from flask_login import UserMixin, login_user
from markdown import markdown
from docx import Document
#import openai
#from openai import OpenAI
from source.legacy.utilities import db_helper, queries
from source.common import config
import pymysql
import bcrypt
import os.path
import markdown
from urllib.parse import urlparse
from pymongo import MongoClient
from bson.objectid import ObjectId
from werkzeug.utils import secure_filename
import jwt

# For this example, we use dictionary to represent users.
# Replace this with DB based user management system

#orgs = []
#users = []
#orgname = ""
userdocs_path = os.path.join('data','user_docs')
os.makedirs(userdocs_path, exist_ok=True)
#org_filename = os.path.join(userdocs_path,'orgs.txt')


#with open(org_filename, "r") as file:
#    orgs = json.load(file)

#openai.api_key = config.openai_api_key

JSON_FOLDER = os.path.join('data', 'admin','prompts')

# Get the MongoDB connection URI from the environment variable
mongo_uri = config.MONGO_URI

# Connect to MongoDB
mongo_client = MongoClient(mongo_uri)

class User(UserMixin):
    pass

def clean_response_for_html(response):
   response = response.replace('\u20b9', 'INR')
   response =  re.sub(r'```python', '<pre><code>', response, count=1)
   response =  re.sub(r'```', '</code></pre>', response, count=1)
   return response


def ensure_directory_exists(path: str):
    """Ensure the provided directory path exists."""
    os.makedirs(path, exist_ok=True)

def find_org_name(username):
    """
    Retrieve the organization name for a given username.

    Args:
        username (str): The username for which to find the organization name.

    Returns:
        str: The organization name if found, an empty string otherwise.
    """
    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    # Use a SQL JOIN to retrieve the organization name for the given username
    sql = """
    SELECT o.org_name
    FROM organization o
    JOIN user_organization uo ON o.org_id = uo.org_id
    JOIN users u ON u.id = uo.user_id
    WHERE u.username = %s
    """
    cursor.execute(sql, (username,))
    rows = cursor.fetchall()
    result = [row["org_name"] for row in rows]
    cursor.close()
    connection.close()

    if result:
        return result
    return []


def chatname_edits(chat):
    # Replace all characters after the 10th with '...'
    if len(chat) > 30:
        chat = chat[:30] + '...'

    # Replace all '_' with space
    chat = chat.replace('_', ' ')
    return chat

def get_old_chats():
    # List all the files in the folder specified by 'folder_path'
    folder_path = os.path.join(userdocs_path,session.get('orgname', ''),'users',session.get('username', ''))
    # Ensure the directory exists
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
        # Sort the files by modification date, latest first
        old_chats_temp = sorted(files, key=lambda x: os.stat(os.path.join(folder_path, x)).st_mtime, reverse=True)
        #remove underscore, shorten
        old_chats = []
        for chat in old_chats_temp:
            if not chat == 'currentchat':
                # Store the original chat name
                original_chat = chat
                chat = chatname_edits(chat)

                old_chats.append([original_chat, chat])

    else:
        old_chats = []
    return old_chats

def get_old_chats_index():
    orgname = session.get('orgname', '')
    username = session.get('username', '')

    # Select the database and collection
    #db = mongo_client['pw-api']  # Database name
    db_chatlog = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['chatlog']  # Collection name

    # Define the filter criteria
    filter_criteria = {
        'username': username,
        'orgname' : orgname
    }
    chatlog_documents = db_chatlog.find(filter_criteria)
    #extract list of all chats with title and id
    old_chats = [[doc['_id'], doc['title']] for doc in chatlog_documents]
    return old_chats

def get_old_chat_logs(folder_path):
    # Step 1: Open the file named currentfile within folder_path to read the name of the log file
    with open(os.path.join(folder_path, "currentchat"), "r") as file:
        log_file_name = file.readline().strip()

    if log_file_name == 'new_chat':
        return [], "", "", "", ""

    # Step 2: Open the log file and read its entire content
    with open(os.path.join(folder_path, log_file_name), "r") as log_file:
        lines = log_file.readlines()

    if not lines:
        return [], 'docchat', '', 'misc', ""

    # Process the first line for metadata
    metadata = json.loads(lines[0].strip())
    type = metadata.get('type', 'docchat')
    filename = metadata.get('file', '')
    collection = metadata.get('collection', '')

    chat_logs = []
    metadata_array = []
    current_response = ""

    # Iterate through the remaining lines
    for line in lines[1:]:
        line = line.strip()

        # Check if the line is a new JSON object (i.e., starts with '{')
        if line.startswith('{'):
            if current_response:
                # Process the previous response
                data = json.loads(current_response)
                question = markdown.markdown(data.get("question", ""))
                answer = markdown.markdown(data.get("response", ""))
                id = data.get("id", "")
                chat_logs.append([id, question, answer])
                metadata_str = json.dumps(data.get("metadata", ""))
                metadata_array.append(metadata_str)
                current_response = ""

            current_response = line
        else:
            # Append continuation of the response to the current_response
            current_response += " " + line

    # Process the last response if it exists
    if current_response:
        data = json.loads(current_response)
        question = markdown.markdown(data.get("question", ""))
        answer = markdown.markdown(data.get("response", ""))
        id = data.get("id", "")
        chat_logs.append([id, question, answer])
        metadata_str = json.dumps(data.get("metadata", ""))
        metadata_array.append(metadata_str)


    return chat_logs, type, filename, collection, metadata_array

def get_old_chat_logs_new(latest_chat_thread_id):
    orgname = session.get('orgname', '')
    username = session.get('username', '')

    latest_chatlog_document = None
    # Select the database and collection
    #db = mongo_client['pw-api']  # Database name
    db_chatlog = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['chatlog']  # Collection name
    db_users = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['users']  # Collection name

    # Define the filter criteria
    filter_criteria = {
        'username': username,
        'orgname' : orgname
    }

    # Query the collections
    user_document = db_users.find_one(filter_criteria)

    # Check if user_document exists, if not insert a new user document
    if user_document is None:
        # Here, decide what default values you want to set for a new user
        new_user_document = {
            'username': username,
            'orgname': orgname,
            'latestchatthreadid': "new_chat"  # Assuming default latest chat thread id is an empty string
        }
        db_users.insert_one(new_user_document)
        user_document = new_user_document
        return [], [], None, None, []

    # Set the latest chat thread id from querystring or extract the latest chat thread ID from DB
    if (latest_chat_thread_id is None):
        if user_document['latestchatthreadid'] != "new_chat":
            latest_chat_thread_id = user_document['latestchatthreadid']
    else:
        # Update the document
        update_result = db_users.update_one(
            filter_criteria,
            {'$set': {'latestchatthreadid': latest_chat_thread_id}}
        )
    print(latest_chat_thread_id)
    if latest_chat_thread_id != '':
        latest_chat_thread_id_obj = ObjectId(latest_chat_thread_id)

        #now get latest chatlog thread
        filter_criteria = {
            "_id": latest_chat_thread_id_obj
        }
        latest_chatlog_document = db_chatlog.find_one(filter_criteria)
    if latest_chatlog_document is None:
        # Handle the case where there is no chatlog document found
        return [], [], None, None, []

    module = latest_chatlog_document["type"]
    collection = latest_chatlog_document["collection"]

    # Initialize an empty array to store the chatlogs
    old_chat_logs = []
    metadata_array_str = []
    chat_ids = []

    for chatlog in latest_chatlog_document['chatlogs']:
        # Extract id, question, and answer from each chatlog and append to the array
        old_chat_log = {"role": chatlog.get('role'), "content": markdown.markdown(chatlog.get('content'))}
        old_chat_logs.append(old_chat_log)
        metadata_array_str.append(chatlog.get("metadata"))
        chat_ids.append(chatlog.get('id'))

    return old_chat_logs, chat_ids, module, collection, metadata_array_str

def check_latest_chat_thread_id(latest_chat_thread_id):

    orgname = session.get('orgname', '')
    username = session.get('username', '')

    # Select the database and collection
    #db = mongo_client['pw-api']  # Database name
    db_chatlog = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['chatlog']  # Collection name


    try:
        # Convert the string to an ObjectId
        object_id = ObjectId(latest_chat_thread_id)

        # Define the filter criteria
        filter_criteria = {
            "_id": object_id,
            "username": username,
            "orgname": orgname
        }

        # Check if a document with the given _id, username, and orgname exists
        chatlog_document = db_chatlog.find_one(filter_criteria)

        # Return True if the document exists, False otherwise
        return chatlog_document is not None
    except Exception as e:
        print(f"An error occurred: {e}")
        return False

def set_latestchatthreadid(latest_chat_thread_id, username, orgname):
    # Select the database and collection
    #db = mongo_client['pw-api']  # Database name
    db_user = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['users']  # Collection name

    # Define the filter criteria and the update operation
    filter_criteria = {'username': username, 'orgname': orgname}
    update_data = {'$set': {'latestchatthreadid': latest_chat_thread_id}}

    try:
        # Perform the update
        update_result = db_user.update_one(filter_criteria, update_data)

        # Check if the update was successful
        return update_result.modified_count > 0
    except Exception as e:
        print(f"An error occurred: {e}")
        return False

from flask import session

def get_latestchatthreadid():
    # Select the database and collection
    #db = mongo_client['pw-api']  # Database name
    db_user = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['users']  # Collection name


    # Retrieve username and orgname from the session
    username = session.get('username')
    orgname = session.get('orgname')

    # Define the filter criteria
    filter_criteria = {'username': username, 'orgname': orgname}

    try:
        # Perform the query
        user_document = db_user.find_one(filter_criteria)

        # Check if the document was found and return the latestchatthreadid
        if user_document and 'latestchatthreadid' in user_document:
            print(f"inside get_latestchatthreadid: {user_document['latestchatthreadid']}")
            return user_document['latestchatthreadid']
    except Exception as e:
        print(f"An error occurred: {e}")
    return "new_chat"

def insert_chatlog(latest_chat_thread_id, chatlog_id, role, content, metadata):

    # Check if latest_chat_thread_id is valid
    if not latest_chat_thread_id:
        print("Invalid latest_chat_thread_id provided")
        return False

    # Select the database and collection
    #db = mongo_client['pw-api']  # Database name
    db_chatlog = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['chatlog']  # Collection name


    # Convert the string to an ObjectId
    object_id = ObjectId(latest_chat_thread_id)

    # Define the new chatlog entry
    new_chatlog = {
        "id": chatlog_id,
        "role": role,
        "content": content,
        "metadata": metadata
    }

    try:
        # Update the document by pushing a new chatlog into the 'chatlogs' array
        update_result = db_chatlog.update_one(
            {"_id": object_id},
            {"$push": {"chatlogs": new_chatlog}}
        )

        # Check if the update was successful
        return update_result.modified_count > 0
    except Exception as e:
        print(f"An error occurred: {e}")
        return False

def insert_chatthread(title, module, collection, chatlog_id, role, content, metadata, username, orgname):

    # Select the database and collection
    #db = mongo_client['pw-api']  # Database name
    db_chatlog = mongo_client.get_database(os.getenv('MONGO_DATABASE'))['chatlog']  # Collection name


    if (content != ''):
        # Define the new chatlog entry
        new_chatthread = {
            "username": username,
            "orgname": orgname,
            "title": title,
            "type": module,
            "collection": collection,
            "chatlogs": [{
                "id": chatlog_id,
                "role": role,
                "content": content,
                "metadata": metadata
            }]
        }
    else:
        new_chatthread = {
            "username": username,
            "orgname": orgname,
            "title": title,
            "type": module,
            "collection": collection,
            "chatlogs": []
        }

    try:
        # Update the document by pushing a new chatlog into the 'chatlogs' array
        insert_result = db_chatlog.insert_one(new_chatthread)
        latest_chat_thread_id = insert_result.inserted_id
        set_latestchatthreadid(latest_chat_thread_id, username, orgname)

        # Check if the update was successful
        return str(latest_chat_thread_id)
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def determine_list_type(p):
    """Determine the type of list for the given paragraph using heuristics."""
    #logging.debug(p.style.name)
    if p.style.name == 'List Paragraph':
        #if p.text.startswith(('•', '-', '*')):
        return 'unordered'
        #if any(char.isdigit() for char in p.text[:3]):  # Check the first few characters for digits
            #return 'ordered'
    return None



def docx_to_html(docx_file, image_folder):
    doc = Document(docx_file)
    html = []

    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        list_type = determine_list_type(paragraph)
        #logging.debug(list_type)
        if style.startswith('Heading'):
            level = style.split()[-1]
            html.append(f'<h{level}>')
            for run in paragraph.runs:
                if run.bold:
                    html.append('<b>')
                if run.italic:
                    html.append('<i>')
                if run.underline:
                    html.append('<u>')
                html.append(run.text)
                if run.underline:
                    html.append('</u>')
                if run.italic:
                    html.append('</i>')
                if run.bold:
                    html.append('</b>')
            html.append(f'</h{level}>')
        elif list_type == 'ordered':
            html.append('<ol>')
            html.append('<li>')
            for run in paragraph.runs:
                html.append(run.text)
            html.append('</li>')
            html.append('</ol>')
        elif list_type == 'unordered':
            html.append('<ul>')
            html.append('<li>')
            for run in paragraph.runs:
                html.append(run.text)
            html.append('</li>')
            html.append('</ul>')
        else:
            html.append('<p>')
            for run in paragraph.runs:
                #logging.debug(run.text)
                if run.bold:
                    html.append('<b>')
                if run.italic:
                    html.append('<i>')
                if run.underline:
                    html.append('<u>')
                # Check if the run is part of a hyperlink
                is_hyperlink = False
                url = None

                # List child elements
                for child in run._r:
                    print(f"Element: {child.tag}")

                    # List attributes of each child element
                    for attr_name, attr_value in child.attrib.items():
                        print(f"  Attribute: {attr_name} = {attr_value}")

                print(f"Run text: {run.text}")
                # Scenario 1: Direct hyperlinks
                hyperlink_elem = run._r.xpath('.//w:hyperlink')
                if hyperlink_elem:
                    rel_id = hyperlink_elem[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    if rel_id in doc.part.rels:
                        is_hyperlink = True
                        url = doc.part.rels[rel_id].target_ref

                # Scenario 2: Hyperlinks represented using fields
                if not is_hyperlink and 'HYPERLINK' in run.text:
                    is_hyperlink = True
                    # Extracting URL from the field's instruction text
                    url = run.text.split('"')[1] if '"' in run.text else None

                if is_hyperlink and url:
                    #logging.debug("url is: " + url)
                    html.append(f'<a href="{url}">')

                html.append(run.text)

                if is_hyperlink:
                    html.append('</a>')

                if run.underline:
                    html.append('</u>')
                if run.italic:
                    html.append('</i>')
                if run.bold:
                    html.append('</b>')
            html.append('</p>')
        """ elif paragraph.style.name == 'List Number':
            html.append('<ol>')
            for run in paragraph.runs:
                html.append('<li>')
                html.append(run.text)
                html.append('</li>')
            html.append('</ol>')
        elif paragraph.style.name == 'List Bullet':
            html.append('<ul>')
            for run in paragraph.runs:
                html.append('<li>')
                html.append(run.text)
                html.append('</li>')
            html.append('</ul>') """

    for table in doc.tables:
        html.append('<table border="1">')
        for row in table.rows:
            html.append('<tr>')
            for cell in row.cells:
                html.append('<td>')
                for paragraph in cell.paragraphs:
                    html.append(paragraph.text)
                html.append('</td>')
            html.append('</tr>')
        html.append('</table>')

    for rel in doc.part.rels.values():
        """ if "image" in rel.reltype:
            img_path = rel.target_ref
            img_name = os.path.basename(img_path)
            doc.part.get_rel(rel.rId).target_part.blob.save(os.path.join(image_folder, img_name))
            html.append(f'<img src="{img_name}" alt="Image"/>') """

    return ''.join(html)

def remove_non_ascii(s):
    return "".join(c for c in s if ord(c)<128)

def get_files_from_category(json_filelist, folder_name):
    # Parse the JSON string into a list of dictionaries
    filelist = json.loads(json_filelist)

    # Filter filenames where the 'category' key matches the folder_name
    matching_files = [item['filename'] for item in filelist if item['category'] == folder_name]

    return matching_files




def log_to_db(username, user_content, system_content, timestamp_start, timestamp_end, response):
    connection = get_connection()
    cursor = connection.cursor()

    # Insert log entry into the database
    query = (
        "INSERT INTO log_ai_calls (username, user_content, system_content, timestamp_start, timestamp_end, response) "
        "VALUES (%s, %s, %s, %s, %s, %s)"
    )
    cursor.execute(query, (username, user_content, system_content, timestamp_start, timestamp_end, response))
    connection.commit()

    cursor.close()
    connection.close()


def load_prompts(json_filename):
    prompts = []
    prompt_filepath = os.path.join(JSON_FOLDER, json_filename)
    with open(prompt_filepath, 'r') as prompt_file:
        prompts = json.load(prompt_file)
    return prompts

def load_prompts_from_db(prompt_type):
    """
    Create a 2D associative array from a database table based on the prompt_type.

    :param prompt_type: The prompt type to filter the table.
    :return: 2D associative array where the first index is the value of prompt_name and the second index contains
             values from prompt_text, name_label, and description fields.
    """
    associative_array = {}

    # Connect to the database
    connection = get_connection()

    try:
        with connection.cursor() as cursor:
            # SQL query to select the required fields from the table based on prompt_type
            sql = "SELECT prompt_name, prompt_text, name_label, description FROM prompts WHERE prompt_type = %s"
            cursor.execute(sql, (prompt_type,))

            # Fetch all the rows
            rows = cursor.fetchall()

            # Process each row
            for row in rows:
                prompt_name, prompt_text, name_label, description = row
                associative_array[prompt_name] = {
                    'value': prompt_text,
                    'label': name_label,
                    'description': description
                }
    finally:
        connection.close()

    return associative_array



def is_user_admin(username):
    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # Query the users table to get the is_admin value for the given username
            sql = "SELECT is_admin FROM users WHERE username=%s"
            cursor.execute(sql, (username,))
            result = cursor.fetchone()

            # Check if the result exists and if is_admin is 1
            if result and result[0] == 1:
                return True
    finally:
        # Close the database connection
        connection.close()

    return False


# Connection function
def get_connection():

    # Parse the URI
    parsed_uri = urlparse(config.DATABASE_URI)
    return pymysql.connect(host=parsed_uri.hostname,
                             user=parsed_uri.username,
                             password=parsed_uri.password,
                             database=parsed_uri.path.lstrip('/'),
                             port=parsed_uri.port)

# USERS TABLE CRUD

def create_user(username, plain_password, email):  # Add more fields as needed
    # Hash the password
    hashed_password = bcrypt.hashpw(plain_password.encode('utf-8'), bcrypt.gensalt())

    connection = get_connection()
    cursor = connection.cursor()
    sql = "INSERT INTO users (username, password_hash, email) VALUES (%s, %s, %s)"
    cursor.execute(sql, (username, hashed_password, email))  # Add more fields as needed
    connection.commit()
    cursor.close()
    connection.close()

def get_user(user_id):
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "SELECT * FROM users WHERE id = %s"
        cursor.execute(sql, (user_id,))
        result = cursor.fetchone()
    connection.close()
    return result

def get_users_as_json():
    """
    Fetch users and their hashed passwords from the database
    and format it as a JSON list of dictionaries.

    Returns:
        list[dict]: List of dictionaries with usernames and their hashed passwords.
    """
    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    sql = "SELECT username, password_hash FROM users"
    cursor.execute(sql)
    results = cursor.fetchall()

    cursor.close()
    connection.close()

    # Format the data
    users_list = [{user['username']: {"pw": user['password_hash']}} for user in results]

    return users_list


def update_user(user_id, new_username):  # Add more fields as needed
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "UPDATE users SET username = %s WHERE id = %s"
        cursor.execute(sql, (new_username, user_id))  # Add more fields as needed
    connection.commit()
    connection.close()

def delete_user(user_id):
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "DELETE FROM users WHERE id = %s"
        cursor.execute(sql, (user_id,))
    connection.commit()
    connection.close()

def validate_user(username, plain_password):
    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)
    sql = "SELECT password_hash FROM users WHERE username = %s"
    cursor.execute(sql, (username,))
    result = cursor.fetchone()
    cursor.close()
    connection.close()

    if result:
        stored_hashed_password = result['password_hash']
        return bcrypt.checkpw(plain_password.encode('utf-8'), stored_hashed_password.encode('utf-8'))
    return False



# ORGANIZATION TABLE CRUD
def create_organization(org_name):
    """
    Add a new organization to the database.

    Args:
        org_name (str): Name of the organization to add.

    Returns:
        int: ID of the newly created organization.
    """
    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    sql = "INSERT INTO organization (org_name) VALUES (%s)"
    cursor.execute(sql, (org_name,))

    # Get the ID of the newly created organization
    new_org_id = cursor.lastrowid

    #add default collection for org
    create_default_collection_org(new_org_id, cursor)

    connection.commit()
    cursor.close()
    connection.close()

    #now add the folder inside the Uploads Dir for this org
    pathlib.Path(os.path.join(config.ORGDIR_PATH, org_name)).mkdir(exist_ok=True)
    if not os.path.exists(config.ORGFILES_BASE_DIR):
        pathlib.Path(config.ORGFILES_BASE_DIR).mkdir(exist_ok=True)
    if not os.path.exists(os.path.join(config.ORGFILES_BASE_DIR, org_name)):
        pathlib.Path(os.path.join(config.ORGFILES_BASE_DIR, org_name)).mkdir(exist_ok=True)

    return new_org_id

def create_default_collection_org(new_org_id, cursor):
    sql = "INSERT INTO collections (collection_name, metatadata_prompt_prelude	) VALUES ('NotInAnyCollection', 'Please return strictly formatted JSON for the text in Context with following fields:\n')"
    cursor.execute(sql)

    # Get the ID of the newly created collection
    new_collection_id = cursor.lastrowid

    sql = "INSERT INTO organization_collections (org_id, collection_id) VALUES (%s , %s)"
    cursor.execute(sql, (new_org_id, new_collection_id))

    return

# USER_ORGANIZATION TABLE CRUD

def add_user_to_org(user_id, org_id, role='MEMBER'):
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "INSERT INTO user_organization (user_id, org_id, role) VALUES (%s, %s, %s)"
        cursor.execute(sql, (user_id, org_id, role))
    connection.commit()
    connection.close()

def get_orgs_of_user(user_id):
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "SELECT org_id FROM user_organization WHERE user_id = %s"
        cursor.execute(sql, (user_id,))
        result = cursor.fetchall()
    connection.close()
    return result

def get_users_of_org(org_id):
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "SELECT user_id FROM user_organization WHERE org_id = %s"
        cursor.execute(sql, (org_id,))
        result = cursor.fetchall()
    connection.close()
    return result

def remove_user_from_org(user_id, org_id):
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "DELETE FROM user_organization WHERE user_id = %s AND org_id = %s"
        cursor.execute(sql, (user_id, org_id))
    connection.commit()
    connection.close()

def get_org_id_by_name(org_name):
    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    sql = "SELECT org_id FROM organization WHERE org_name = %s"
    cursor.execute(sql, (org_name,))
    result = cursor.fetchone()

    cursor.close()
    connection.close()

    if result:
        return result['org_id']
    return None

def get_all_organizations():

    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    sql = "SELECT * FROM organization"
    cursor.execute(sql)
    results = cursor.fetchall()

    cursor.close()
    connection.close()

    return results

def get_user_id_by_username(username):

    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    sql = "SELECT id FROM users WHERE username = %s"
    cursor.execute(sql, (username,))
    result = cursor.fetchone()

    cursor.close()
    connection.close()

    if result:
        return result['id']
    return None

def get_user_by_username(username):
    """
    Check if a user with the given username exists in the database.

    Args:
        username (str): The username to check.

    Returns:
        bool: True if the user exists, False otherwise.
    """
    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    sql = "SELECT username FROM users WHERE username = %s"
    cursor.execute(sql, (username,))
    result = cursor.fetchone()

    cursor.close()
    connection.close()

    return bool(result)  # Returns True if user exists, False otherwise


def get_organizations_with_users():
    connection = get_connection()
    cursor = connection.cursor(pymysql.cursors.DictCursor)

    # Retrieve all organizations
    org_sql = "SELECT org_id, org_name FROM organization"
    cursor.execute(org_sql)
    orgs = cursor.fetchall()

    # For each organization, retrieve associated users
    for org in orgs:
        user_sql = """
        SELECT u.username AS name
        FROM users u
        JOIN user_organization uo ON u.id = uo.user_id
        WHERE uo.org_id = %s
        """
        cursor.execute(user_sql, (org['org_id'],))
        users = cursor.fetchall()
        org['users'] = users
        del org['org_id']  # Remove org_id as it's not needed in the final result

    cursor.close()
    connection.close()

    return orgs

#some initialisations use the functions above and so given here
""" users_filename = os.path.join(userdocs_path,'users.txt')
with open(users_filename, "r") as file:
    users = json.load(file) """
users = get_users_as_json()


def find_org_id(org_name):
    connection = get_connection()
    cursor = connection.cursor()

    query = "SELECT org_id FROM organization WHERE org_name = %s LIMIT 1"
    cursor.execute(query, (org_name,))
    result = cursor.fetchone()

    cursor.close()
    connection.close()

    if result:
        return result[0]
    else:
        return None


def fetch_modules_by_org(org_id):
    connection = get_connection()
    cursor = connection.cursor()

    # You might need to adjust the query based on your exact table and field names
    query = """SELECT m.form_action, m.button_text, m.name
               FROM organization_modules om
               JOIN modules m ON om.module_id = m.id
               WHERE om.org_id = %s"""
    cursor.execute(query, (org_id,))
    rows = cursor.fetchall()

    # Assuming form_action and button_text are columns in the modules table (or adjust to your actual column names)
    modules = [{'form_action': row[0], 'button_text': row[1], 'name': row[2]} for row in rows]

    cursor.close()
    connection.close()
    return modules

def get_org_default_module_function(orgname):
    org_default_module_function = None

    org_id = get_org_id_by_name(orgname)
    #org_default_module_function = generic_query('organization','home_module_function_name', {"org_id": org_id})[0][0]
    query_result = generic_query('organization', 'home_module_function_name', {"org_id": org_id})
    if query_result:
        org_default_module_function = query_result[0][0]
    else:
        org_default_module_function = None  # or a suitable default value
    return org_default_module_function



def get_filelist(orgname, folder):
    base_path = os.path.join('data', 'user_docs', orgname, 'files', 'docchat2')
    folder_path = os.path.join(base_path, 'folders', folder, 'processed')
    filelist = []

    if os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            filelist.append(file)
    return filelist

def update_page_num_processed(cur_dir, filename, pnum):
    index_file_path = os.path.join(cur_dir, 'data', 'unprocessed', 'index.json')

    # Step 1: Open the index file
    try:
        with open(index_file_path, 'r+', encoding='utf-8') as file:
            # Step 2: Load the JSON data from the file
            data = json.load(file)

            # Step 3: Find the item with the matching 'saved_file_path_name'
            item_found = False
            for item in data:
                if item['saved_file_path_name'] == filename:
                    # Step 4: Replace the 'page_num_processed' value
                    item['page_num_processed'] = str(pnum)
                    item_found = True
                    break

            # Check if the item was found and updated
            if not item_found:
                raise ValueError(f"No item found with 'saved_file_path_name' = {filename}")

            # Step 5: Save the changes back to the JSON file
            # Move the pointer to the beginning of the file to overwrite it
            file.seek(0)
            json.dump(data, file, indent=4)
            file.truncate()  # Truncate the file in case new data is smaller than old

    except FileNotFoundError:
        print(f"The file at {index_file_path} was not found.")
    except json.JSONDecodeError:
        print("Error decoding JSON from the file.")
    except ValueError as ve:
        print(ve)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def update_sectionnum_processed(cur_dir, filename, sectionnum):
    index_file_path = os.path.join(cur_dir, 'data', 'unprocessed', 'index.json')

    # Step 1: Open the index file
    try:
        with open(index_file_path, 'r+', encoding='utf-8') as file:
            # Step 2: Load the JSON data from the file
            data = json.load(file)

            # Step 3: Find the item with the matching 'saved_file_path_name'
            item_found = False
            for item in data:
                if item['saved_file_path_name'] == filename:
                    # Step 4: Replace the 'page_num_processed' value
                    item['sectionnum_processed'] = str(sectionnum)
                    item_found = True
                    break

            # Check if the item was found and updated
            if not item_found:
                raise ValueError(f"No item found with 'saved_file_path_name' = {filename}")

            # Step 5: Save the changes back to the JSON file
            # Move the pointer to the beginning of the file to overwrite it
            file.seek(0)
            json.dump(data, file, indent=4)
            file.truncate()  # Truncate the file in case new data is smaller than old

    except FileNotFoundError:
        print(f"The file at {index_file_path} was not found.")
    except json.JSONDecodeError:
        print("Error decoding JSON from the file.")
    except ValueError as ve:
        print(ve)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")


def update_text_structuring_completed(cur_dir, filename, status):
    index_file_path = os.path.join(cur_dir, 'data', 'unprocessed', 'index.json')

    # Step 1: Open the index file
    try:
        with open(index_file_path, 'r+', encoding='utf-8') as file:
            # Step 2: Load the JSON data from the file
            data = json.load(file)

            # Step 3: Find the item with the matching 'saved_file_path_name'
            item_found = False
            for item in data:
                if item['saved_file_path_name'] == filename:
                    # Step 4: Replace the 'page_num_processed' value
                    item['text_structuring_completed'] = status
                    item_found = True
                    break

            # Check if the item was found and updated
            if not item_found:
                raise ValueError(f"No item found with 'saved_file_path_name' = {filename}")

            # Step 5: Save the changes back to the JSON file
            # Move the pointer to the beginning of the file to overwrite it
            file.seek(0)
            json.dump(data, file, indent=4)
            file.truncate()  # Truncate the file in case new data is smaller than old

    except FileNotFoundError:
        print(f"The file at {index_file_path} was not found.")
    except json.JSONDecodeError:
        print("Error decoding JSON from the file.")
    except ValueError as ve:
        print(ve)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def get_text_structuring_status(cur_dir, filename):
    index_file_path = os.path.join(cur_dir, 'data', 'unprocessed', 'index.json')

    # Step 1: Open the index file
    try:
        with open(index_file_path, 'r', encoding='utf-8') as file:
            # Step 2: Load the JSON data from the file
            data = json.load(file)

            # Step 3: Find the item with the matching 'saved_file_path_name'
            for item in data:
                if item['saved_file_path_name'] == filename:
                    # Step 4: Retrieve the 'text_structuring_completed' value
                    return item.get('text_structuring_completed', None)

            # If we reach this point, no matching item was found
            print(f"No item found with 'saved_file_path_name' = {filename}")
            return None

    except FileNotFoundError:
        print(f"The file at {index_file_path} was not found.")
        return None
    except json.JSONDecodeError:
        print("Error decoding JSON from the file.")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None



def safe_copy(source_filename, saved_file_path_name):
    # Check if the destination file already exists
    if os.path.exists(saved_file_path_name):
        print(f"The file {saved_file_path_name} already exists.")
        # Handle the situation: either remove the old file or change the name of the new one
        try:
            os.remove(saved_file_path_name)  # This will remove the file if it exists
        except PermissionError as e:
            print(f"Unable to remove the existing file: {e}")
            return False

    try:
        shutil.copy(source_filename, saved_file_path_name)
        print(f"File copied successfully to {saved_file_path_name}.")
    except PermissionError as e:
        print(f"Permission denied: {e}")
        return False
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return False

    return True

def getfolder_nodecontent(node_content):
    node_content_json = json.loads(node_content)
    folder = "no folder specified"
    if 'file_path' in node_content_json['metadata']:
        parts = node_content_json['metadata']['file_path'].split('\\')
        if len(parts) > 8:
            folder = parts[8]  # 8th part is the folder name
        #example of mode_content file_path: "file_path": "..\\pw-api\\data\\user_docs\\org3\\files\\docchat2\\folders\\general\\unprocessed\\2. SMF - DSC_2023-11-10_17-27.pdf"

    return folder


#function to handle collections and files using the new vectorisation approach - 9-Dec-2023
def get_collections_for_org(orgname):
    collections_list = []
    # Get org_id for the given orgname
    org_id = get_org_id_by_name(orgname)

    # Establish a connection to the database
    connection = get_connection()
    try:

        # Create a cursor object
        with connection.cursor() as cursor:
            # SQL query to fetch collection names and descriptions
            sql = "SELECT collection_name, description FROM collections WHERE collection_id IN (SELECT collection_id FROM organization_collections WHERE org_id=%s)"

            # Execute the query
            cursor.execute(sql, (org_id))

            # Fetch all the records
            rows = cursor.fetchall()

            # Format the results into the desired JSON structure
            collections_list = [{"collection": row[0], "description": row[1]} for row in rows]

            # Convert the list to JSON format
            return collections_list
    except pymysql.MySQLError as e:
        print(f"Error: {e}")
        return collections_list
    finally:
        # Close the connection
        connection.close()

def get_files_and_collections_for_org(orgname):

    org_id = get_org_id_by_name(orgname)
    connection = get_connection()
    try:
        with connection.cursor() as cursor:

            # SQL query to fetch file names and collection names for the given org_id
            sql = """
            SELECT f.file_name, c.collection_name 
            FROM files f
            JOIN files_collections fc ON f.file_id = fc.file_id
            JOIN collections c ON fc.collection_id = c.collection_id
            JOIN organization_collections oc ON c.collection_id = oc.collection_id
            WHERE oc.org_id = %s
            """

            cursor.execute(sql, (org_id,))
            rows = cursor.fetchall()

            return [{"filename": row[0], "collection": row[1]} for row in rows]
    except pymysql.MySQLError as e:
        print(f"Error: {e}")
        return []
    finally:
        connection.close()

def get_files_for_collection_and_org(foldername, orgname):
    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # Get org_id for the given orgname
            org_id = get_org_id_by_name(orgname)

            # SQL query to fetch file names for the given collection name and org_id
            sql = """
            SELECT f.file_name 
            FROM files f
            JOIN files_collections fc ON f.file_id = fc.file_id
            JOIN collections c ON fc.collection_id = c.collection_id
            JOIN organization_collections oc ON c.collection_id = oc.collection_id
            WHERE c.collection_name = %s AND oc.org_id = %s
            """

            cursor.execute(sql, (foldername, org_id))
            rows = cursor.fetchall()

            return [{"filename": row[0], "collection": foldername} for row in rows]
    except pymysql.MySQLError as e:
        print(f"Error: {e}")
        return []
    finally:
        connection.close()

def get_collection_id_by_name(orgname, collection_name):
    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # Updated SQL query to join with organization_collections and organization tables
            sql = """
            SELECT c.collection_id 
            FROM collections c
            JOIN organization_collections oc ON c.collection_id = oc.collection_id
            JOIN organization o ON oc.org_id = o.org_id
            WHERE c.collection_name = %s AND o.org_name = %s
            """
            cursor.execute(sql, (collection_name, orgname))
            result = cursor.fetchone()
            return result[0] if result else None
    finally:
        connection.close()


""" def update_file_collections(orgname, file_names, new_collection_name):
    new_collection_id = get_collection_id_by_name(orgname, new_collection_name)
    if new_collection_id is None:
        raise ValueError("Collection not found")

    # Format the file_names list into a string for SQL query
    formatted_file_names = ', '.join(["%s"] * len(file_names))

    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # Update query using formatted string
            sql = f<put 3 doubles quotes here>
            UPDATE files_collections 
            SET collection_id = %s 
            WHERE file_id IN (
                SELECT file_id 
                FROM files 
                WHERE file_name IN ({formatted_file_names})
            )
            <put 3 doubles quotes here>
            cursor.execute(sql, [new_collection_id] + file_names)
            connection.commit()
    finally:
        connection.close() """

def update_file_collections(orgname, file_names, new_collection_name):
    new_collection_id = get_collection_id_by_name(orgname, new_collection_name)
    if new_collection_id is None:
        raise ValueError("Collection not found")

    org_id = get_org_id_by_name(orgname)
    if org_id is None:
        raise ValueError("Organization not found")

    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # Update query to consider organization
            sql = """
            UPDATE files_collections fc
            INNER JOIN files f ON fc.file_id = f.file_id
            INNER JOIN organization_collections oc ON fc.collection_id = oc.collection_id
            SET fc.collection_id = %s
            WHERE f.file_name IN %s AND oc.org_id = %s
            """

            # Using tuples for the IN clause in the SQL query
            cursor.execute(sql, (new_collection_id, tuple(file_names), org_id))
            connection.commit()
    finally:
        connection.close()

#import pinecone
from pinecone import Pinecone


def update_pinecone_vectors_metadata(selected_files, org_name, new_folder_name):
    #logger = set_debug_logger()
    # Initialize Pinecone client and connect to your index
    """ pinecone.init(api_key=config.pinecone_api_key, environment=config.pinecone_environment)
    index = pinecone.Index(config.pinecone_index_name) """
    pc = Pinecone(api_key=config.PINECONE_API_KEY)
    index = pc.Index(config.PINECONE_INDEX_NAME)

    """ # Query to fetch vectors by filename and orgname
    query_results = index.query([0] * 1536, 
                                filter={"orgname": {"$eq": org_name}, 
                                        "filename": {"$in": selected_files}}, 
                                top_k=10000, 
                                include_metadata=True, include_values=True)

    # Update 'collection' metadata for each matching vector
    updates = []
    #logger.debug("# of Vectors matched: " + str(len(query_results['matches'])))
    for match in query_results['matches']:
        vector_id = match['id']
        metadata_update = match['metadata']
        metadata_update['collection'] = new_folder_name
        updates.append((vector_id, match['values'], metadata_update))

    # Perform batch upsert with the updated metadata
    #logger.debug("# of Vectors updated: " + str(len(updates)))
    if updates:
        response = index.upsert(vectors=updates)
        print("Upsert response:", response) """

    # Query to fetch vectors by filename and orgname
    query_results = index.query(vector=[0] * 1536,
                                filter={"orgname": {"$eq": org_name},
                                        "filename": {"$in": selected_files}},
                                top_k=10000,
                                include_metadata=True, include_values=False)

    # Update 'collection' metadata for each matching vector
    #ids = []
    #logger.debug("# of Vectors matched: " + str(len(query_results['matches'])))
    for match in query_results['matches']:
        #vector_id = match['id']
        #ids.append(vector_id)
        response = index.update(id=match['id'], set_metadata={"collection": new_folder_name})
    #print(f"ids: {updates}")

    # Perform batch upsert with the updated metadata
    #logger.debug("# of Vectors updated: " + str(len(updates)))
    #if updates:

    #print("Upsert response:", response)

    #pinecone.close()

def create_new_collection_with_org(name, description, orgname):
    org_id = get_org_id_by_name(orgname)
    if org_id is None:
        raise ValueError("Organization not found")

    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # Insert into collections table
            sql = "INSERT INTO collections (collection_name, description, metatadata_prompt_prelude) VALUES (%s, %s, %s)"
            cursor.execute(sql, (name, description, 'Please return strictly formatted JSON for the text in Context with following fields:\n'))
            collection_id = cursor.lastrowid

            # Insert into organization_collections table
            sql = "INSERT INTO organization_collections (org_id, collection_id) VALUES (%s, %s)"
            cursor.execute(sql, (org_id, collection_id))
            connection.commit()
    finally:
        connection.close()

def get_all_collections_for_org(orgname):
    org_id = get_org_id_by_name(orgname)
    if org_id is None:
        raise ValueError("Organization not found")

    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            sql = """
            SELECT c.collection_name 
            FROM collections c
            JOIN organization_collections oc ON c.collection_id = oc.collection_id
            WHERE oc.org_id = %s
            """
            cursor.execute(sql, (org_id,))
            rows = cursor.fetchall()
            return [row[0] for row in rows]
    finally:
        connection.close()

#The deletion happens outside of the backend and so not needed. retained for later use, if any
def delete_files_from_system(org_name, filenames):
    folder_path = os.path.join(config.ORGDIR_PATH, org_name)
    for filename in filenames:
        file_path = os.path.join(folder_path, filename)
        if os.path.exists(file_path):
            os.remove(file_path)

#this way deleting files assumes: vectoriser calls this backend with filenames and it only needs to be deleted from the DB. Also filenames will be unique within an Org and may be same between 2 orgs. And all files when added are added by default to NotInAnyCollection collection. So once it is deleted below from collection, then any file not in any collection will be the correct file to delete.
def delete_files_from_db(org_name, filenames):
    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # First, get the org_id for the given org_name
            org_id = get_org_id_by_name(org_name)

            # Delete from files_collections for the specified org_id and filenames
            sql = """
            DELETE fc FROM files_collections fc
            INNER JOIN files f ON fc.file_id = f.file_id
            INNER JOIN organization_collections oc ON fc.collection_id = oc.collection_id
            INNER JOIN organization o ON oc.org_id = o.org_id
            WHERE f.file_name IN %s AND o.org_id = %s
            """
            cursor.execute(sql, [tuple(filenames), org_id])

            # Delete from files where file_name matches and not linked to any other collection
            sql = """
            DELETE FROM files 
            WHERE file_name IN %s AND NOT EXISTS (
                SELECT 1 FROM files_collections WHERE file_id = files.file_id
            )
            """
            cursor.execute(sql, [tuple(filenames)])

            connection.commit()
    finally:
        connection.close()

def add_files_to_db(org_name, filenames):
    logger = logging.getLogger('app')
    # Assuming each file is already in the FTP folder, and we just need to update the database
    org_id = get_org_id_by_name(org_name)
    if org_id is None:
        raise ValueError("Organization not found")

    # Link the file to a default collection for the organization
    default_collection_id = get_default_collection_id_for_org(org_id)

    files_path = os.path.join(config.ORGDIR_PATH, org_name)
    if not os.path.exists(files_path):
        os.makedirs(files_path)

    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            for filename in filenames:

                """ #old code - eventually delete
                # Add file to the files table
                sql = "INSERT INTO files (file_name) VALUES (%s)"
                cursor.execute(sql, (filename,))
                file_id = cursor.lastrowid

                if default_collection_id:
                    sql = "INSERT INTO files_collections (file_id, collection_id) VALUES (%s, %s)"
                    cursor.execute(sql, (file_id, default_collection_id)) """

                #check if the file is already for any of the collections within this org
                sql = """
                    SELECT f.file_name 
                    FROM files f 
                    JOIN collections c ON f.collection_id = c.collection_id 
                    WHERE f.file_name = %s AND c.org_id = %s
                """
                cursor.execute(sql, (filename, org_id))
                result = cursor.fetchone()
                if result:
                    logger.debug(f"File '{filename}' already exists in the database for organization '{org_name}'. Skipping insertion.")
                    continue
                # Save in DB
                current_saved_file_path = os.path.join(str(files_path), str(filename))
                file_id = db_helper.execute_and_get_id(queries.V2_INSERT_FILE, filename, current_saved_file_path, default_collection_id, '')
                logger.debug(f"File '{filename}' inserted into the database for organization '{org_name}' with file_id {file_id}.")

            connection.commit()
    finally:
        connection.close()

def get_default_collection_id_for_org(org_id):
    connection = get_connection()
    try:
        with connection.cursor() as cursor:
            # SQL query to find the collection_id of "NotInAnyCollection" for the given org_id
            sql = """
            SELECT c.collection_id 
            FROM collections c
            WHERE c.collection_name = %s AND c.org_id = %s
            """
            cursor.execute(sql, ("NotInAnyCollection", org_id))
            result = cursor.fetchone()
            return result[0] if result else None
    finally:
        connection.close()

def remove_query_string(url):
    parsed_url = urlparse(url)
    # Reconstruct the URL without query string
    return parsed_url.scheme + "://" + parsed_url.netloc + parsed_url.path

""" def get_prompt(**kwargs):


    # Connect to the database
    connection = get_connection()

    try:
        with connection.cursor() as cursor:
            if (kwargs['type'] != 'metadata'):
                # SQL query to fetch prompt_text
                query = "SELECT prompt_text FROM prompts WHERE prompt_name = %s"
                cursor.execute(query, (kwargs['prompt_name'],))
                # Fetch one record
                result = cursor.fetchone()

                if result:
                    return result[0]
                else:
                    return ""  # or raise an exception, based on your preference                
            else:
                # SQL query to fetch prompt_text
                query = "SELECT metatadata_prompt_prelude, metatadata_prompt FROM collections WHERE collection_id = %s"
                cursor.execute(query, (kwargs['collection_id'],))
                # Fetch one record
                result = cursor.fetchone()

                if result:
                    return result[0]+result[1]
                else:
                    return ""  # or raise an exception, based on your preference


    except Exception as e:
        print(f"An error occurred: {e}")
        return ""
    finally:
        connection.close() """

def find_collection_id(collection_name):
    """Find the collection ID based on the collection name."""
    try:
        conn = get_connection()
        with conn.cursor() as cursor:
            query = "SELECT collection_id FROM collections WHERE collection_name = %s"
            cursor.execute(query, (collection_name,))
            result = cursor.fetchone()
            return result[0] if result else None
    finally:
        conn.close()

def find_file_ids(collection_id):
    """Find all file IDs for a given collection ID."""
    try:
        conn = get_connection()
        with conn.cursor() as cursor:
            query = "SELECT file_id FROM files_collections WHERE collection_id = %s"
            cursor.execute(query, (collection_id,))
            return [row[0] for row in cursor.fetchall()]
    finally:
        conn.close()

def get_metadata_for_files(file_ids):
    """Get metadata for each file ID."""
    try:
        conn = get_connection()
        with conn.cursor() as cursor:
            format_strings = ','.join(['%s'] * len(file_ids))
            query = f"SELECT metadata FROM files_metadata_plus WHERE file_id IN ({format_strings})"
            cursor.execute(query, file_ids)
            return [row[0] for row in cursor.fetchall()]
    finally:
        conn.close()

def get_metadata_prompt(collection_id, preludeNeeded):
#current DB structure is such that metadata prompts are in the Collections table itself. All other prompts are in the prompts table with prompt_name as unique. This might need to be changed later.
    connection = get_connection()
    with connection.cursor() as cursor:
        query = "SELECT metatadata_prompt_prelude, metatadata_prompt FROM collections WHERE collection_id = %s"
        cursor.execute(query, (collection_id,))
        # Fetch one record
        result = cursor.fetchone()
        connection.close()
        if result:
            if preludeNeeded:
                #return result[0]+result[1]
                return (result[0] if result[0] is not None else "") + (result[1] if result[1] is not None else "")
            else:
                return result[1]
        else:
            return ""

def set_metadata_prompt(metadata_prompt, collection_id):
#current DB structure is such that metadata prompts are in the Collections table itself. All other prompts are in the prompts table with prompt_name as unique. This might need to be changed later.
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "UPDATE collections SET metatadata_prompt = %s WHERE collection_id = %s"
        cursor.execute(sql, (metadata_prompt, collection_id))  # Add more fields as needed
        result = cursor.fetchone()
    connection.commit()
    connection.close()
    return True

def get_custom_instructions(orgname, collection):
    collection_id = get_collection_id_by_name(orgname, collection)
    custom_instructions_result = generic_query('collections', 'custom_instructions', {"collection_id": collection_id})
    custom_instructions = ''
    if len(custom_instructions_result) > 0 and len(custom_instructions_result[0]) > 0:
        custom_instructions = custom_instructions_result[0][0]
    return custom_instructions

def set_custom_instructions(custom_instructions, collection_id):
    connection = get_connection()
    with connection.cursor() as cursor:
        sql = "UPDATE collections SET custom_instructions = %s WHERE collection_id = %s"
        cursor.execute(sql, (custom_instructions, collection_id))
        result = cursor.fetchone()
    connection.commit()
    connection.close()
    return True

def generic_query(table, column, filters=None):
    """
    Execute a query on a specified table and column, with optional filters.

    :param table: Name of the table to query.
    :param column: Name of the column to retrieve.
    :param filters: Optional dictionary of filters (column: value).
    :return: Query result set.
    """
    # Construct the basic query
    query = f"SELECT {column} FROM {table}"

    # Add filters to the query if provided
    if filters:
        filter_clauses = [f"{key} = %s" for key in filters]
        query += " WHERE " + " AND ".join(filter_clauses)

    # Execute the query
    results = []
    connection = get_connection()  # Assume this function is already defined
    try:
        with connection.cursor() as cursor:
            cursor.execute(query, tuple(filters.values()) if filters else ())
            results = cursor.fetchall()
    finally:
        connection.close()

    return results

#############################################################################
#API functions follow
#############################################################################

def api_v1_login_for_ssr():
    #print(session['login_stage'])
    session.clear()
    auth_methods = []
    if hasattr(config, 'auth_methods'):
        auth_methods = config.AUTH_METHODS_ALLOWED
        #print(auth_methods)

    #if session.get('login_stage') != 'user_validated':
    print("stage 1")
    username = request.form['username']
    password = request.form['password']

    if validate_user(username, password):
        session['username'] = username
        session['newchat'] = '0'

        print("stage 2")
        orgs = find_org_name(username)
        session['login_stage'] = 'user_validated'
        if len(orgs)== 0:
            print("stage 3")
            session['login_stage'] = 'user_not_validated'
            return json.dumps({"render" : "login.html", "redirect" : "login", "auth_methods": auth_methods, "message": "No Org assigned to User. Contact Site Admin.", "orgs": None})
        elif len(orgs) == 1:
            print("stage 4")
            user = User()
            user.id = username
            login_user(user)

            session['orgname'] = orgs[0]
            # Fetch modules for the organization
            org_id = find_org_id(orgs[0])
            modules = fetch_modules_by_org(org_id)
            session['modules'] = modules  # Storing modules in session for easy access in other routes
            org_default_module_function = get_org_default_module_function(session['orgname'])
            if (org_default_module_function is None):
                return json.dumps({"redirect" : "docchat_home", "render" : "", "auth_methods": [], "message": "", "orgs": None})
            else:
                return json.dumps({"redirect" : org_default_module_function, "render" : "", "auth_methods": [], "message": "", "orgs": None})
        else:
            print("stage 5")
            return json.dumps({"render" : "login.html", "redirect" : "login", "auth_methods": auth_methods, "orgs": orgs, "message": ""})
        #return json.dumps({"render" : "login.html", "redirect" : "login", "message": "Error in login", "auth_methods": [], "orgs": None})


def api_v1_login():
    try:
        username = request.json['username']
        password = request.json['password']
        is_user_valid = validate_user(username, password)
        if is_user_valid:
            encoded_jwt = jwt.encode({'username': username}, config.JWT_SECRET, algorithm="HS256")
            return encoded_jwt
        else:
            raise RuntimeError("Invalid user credentials!")
    except Exception as e:
        print(e)
        raise e

def get_all_org_for_user():
    try:
        orgs = find_org_by_user(request.context['username'])
        if len(orgs) > 0:
            return orgs
        else:
            return []
    except Exception as e:
        print(e)
        raise e


def find_org_by_user(username):
    try:
        connection = get_connection()
        cursor = connection.cursor(pymysql.cursors.DictCursor)

        # Use a SQL JOIN to retrieve the organization name and id for the given username
        sql = """
            SELECT o.org_id, o.org_name
            FROM organization o
            JOIN user_organization uo ON o.org_id = uo.org_id
            JOIN users u ON u.id = uo.user_id
            WHERE u.username = %s
            """
        cursor.execute(sql, (username,))
        rows = cursor.fetchall()
        result = [{'org_id': row['org_id'], 'org_name': row["org_name"]} for row in rows]
        cursor.close()
        connection.close()

        if result:
            return result
        return []
    except Exception as e:
        print(e)
        raise e


def api_v1_select_org():
#if session.get('login_stage') == 'user_validated':
    print("stage 6")
    user = User()
    user.id = session['username']
    login_user(user)

    orgname = request.form.get("orgname")
    print(orgname)
    session['orgname'] = orgname
    # Fetch modules for the organization
    org_id = find_org_id(orgname)
    modules = fetch_modules_by_org(org_id)
    session['modules'] = modules  # Storing modules in session for easy access in other routes
    print("stage 7")
    org_default_module_function = get_org_default_module_function(orgname)
    print(org_default_module_function)
    if (org_default_module_function is None):
        return json.dumps({"redirect" : "docchat_home", "render" : "", "auth_methods": [], "orgs": None, "message": ""})
    else:
        return json.dumps({"redirect" : org_default_module_function, "render" : "", "auth_methods": [], "orgs": None, "message": ""})

def api_v1_get_file_list():
    orgname = session.get('orgname', '')

    # Get the list of files and their corresponding collections
    output = get_files_and_collections_for_org(orgname)

    return {"result": output, "orgname": orgname}

def api_v1_get_file(filename):
    # Original PDF path
    orgname = session.get("orgname", "")
    original_pdf_path = os.path.join(config.ORGDIR_PATH, orgname, config.ORGDIR_SUB_PATH, filename)

    # Generate a UUID for the subfolder
    subfolder_name = str(uuid.uuid4())

    # Create the subfolder if it doesn't exist
    api_source_dir = os.path.join('source', 'api')
    tmpfiles_path = os.path.join('static', 'tmpfiles')
    subfolder_name = os.path.join(api_source_dir, tmpfiles_path, subfolder_name)
    if not os.path.exists(subfolder_name):
        os.makedirs(subfolder_name)

    # Define the destination path for the PDF
    filename = secure_filename(filename)
    destination_pdf_path = os.path.join(subfolder_name, filename)

    # Copy the PDF to the new location
    shutil.copyfile(original_pdf_path, destination_pdf_path)

    # Generate a URL to access the PDF
    pdf_internal_path = destination_pdf_path
    file_path = pathlib.Path(pdf_internal_path)
    immediate_folder = file_path.parent.name
    pdf_url = os.path.join(tmpfiles_path, immediate_folder, filename)

    return {'internal_path': pdf_internal_path, 'url': pdf_url}
